import re
import os
import zipfile
import xml.etree.ElementTree as ET
from docxtpl import DocxTemplate
from docx import Document

REVERENSI_DIR = os.path.join(os.path.dirname(__file__), 'reverensi')
TEMPLATE_BA_PATH = os.path.join(REVERENSI_DIR, 'BERITA_ACARA_placeholder (1).docx')
TEMPLATE_DH_PATH = os.path.join(REVERENSI_DIR, 'DAFTAR_HADIR_TEMPLATE.docx')

def parse_excel_bapu(file_path):
    """
    Parses BAPU Excel file exported from UPA Bahasa system.
    Extracts Sesi string, Ruangan, and Student table data.
    """
    ss_map = {}
    students = []
    sesi_raw = ""
    ruangan_raw = "Lab UPA Bahasa"
    
    with zipfile.ZipFile(file_path, 'r') as z:
        if 'xl/sharedStrings.xml' in z.namelist():
            ss_tree = ET.fromstring(z.read('xl/sharedStrings.xml'))
            idx = 0
            for elem in ss_tree.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}si'):
                t_text = ''.join([t.text for t in elem.iter() if t.tag.endswith('t') and t.text])
                ss_map[str(idx)] = t_text
                idx += 1
                
        sheet_tree = ET.fromstring(z.read('xl/worksheets/sheet1.xml'))
        rows_data = []
        for row in list(sheet_tree.iter('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row')):
            row_vals = []
            for c in row.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c'):
                t = c.attrib.get('t')
                v = c.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v')
                val = v.text if v is not None else ''
                if t == 's' and val in ss_map:
                    val = ss_map[val]
                row_vals.append(val.strip())
            rows_data.append(row_vals)

    # Process rows
    for r in rows_data:
        if not r:
            continue
        first_col = r[0]
        if first_col.startswith('Sesi'):
            sesi_raw = first_col.replace('Sesi', '').replace(':', '').strip()
        elif first_col.startswith('Ruangan'):
            ruangan_raw = first_col.replace('Ruangan', '').replace(':', '').strip()
        elif first_col.isdigit():
            no_val = int(first_col)
            nama_val = r[1] if len(r) > 1 else ''
            email_val = r[2] if len(r) > 2 else ''
            nomor_peserta = r[3] if len(r) > 3 else ''
            npm_val = r[4] if len(r) > 4 else ''
            hadir_val = r[5] if len(r) > 5 else ''
            keterangan_val = r[6] if len(r) > 6 else ''
            
            if npm_val == '-' or not npm_val:
                email_prefix = email_val.split('@')[0] if email_val else ''
                if email_prefix.isdigit():
                    npm_val = email_prefix
                else:
                    npm_val = '-'
            
            students.append({
                'id': str(no_val),
                'no': no_val,
                'nama': nama_val,
                'email': email_val,
                'nomor_peserta': nomor_peserta,
                'npm': npm_val,
                'hadir': 'YA' if hadir_val.upper() in ['YA', 'HADIR', 'TRUE', '1'] else 'TIDAK',
                'keterangan': keterangan_val
            })

    # Parse Sesi string: e.g. "Rabu, 05 Agustus 2026 - (pukul 13:30 - 16:30 WIB)"
    hari = "Senin"
    tanggal = "01 Juni 2026"
    jam_mulai = "09.00"
    jam_selesai = "12.00"
    
    if sesi_raw:
        m = re.match(r'^\s*([^,]+),\s*([^\-]+)\s*-\s*\((?:pukul\s*)?([^\)]+)\)', sesi_raw)
        if m:
            hari = m.group(1).strip()
            tanggal = m.group(2).strip()
            waktu_str = m.group(3).replace('WIB', '').strip()
            t_parts = waktu_str.split('-')
            jam_mulai = re.sub(r'\b(\d{2})(\d{2})\b', r'\1.\2', t_parts[0].strip().replace(':', '.'))
            if len(t_parts) > 1:
                jam_selesai = re.sub(r'\b(\d{2})(\d{2})\b', r'\1.\2', t_parts[1].strip().replace(':', '.'))
        else:
            parts = sesi_raw.split(',')
            if len(parts) >= 2:
                hari = parts[0].strip()
                rest = ','.join(parts[1:]).strip()
                t_parts = rest.split('-')
                tanggal = t_parts[0].strip()

    waktu_combined = f"{jam_mulai} - {jam_selesai}"

    return {
        'sesi_raw': sesi_raw,
        'hari': hari,
        'tanggal': tanggal,
        'jam_mulai': jam_mulai,
        'jam_selesai': jam_selesai,
        'waktu': waktu_combined,
        'ruangan': ruangan_raw or "Lab UPA Bahasa",
        'jenis_kegiatan': "ONLINE",
        'students': students
    }

def generate_berita_acara_docx(data, output_path, logo_path=None):
    """
    Generates Berita Acara Word Document (.docx) using reverensi template.
    """
    tpl = DocxTemplate(TEMPLATE_BA_PATH)
    
    total_p = data.get('total_peserta', len(data.get('students', [])))
    hadir_p = data.get('jumlah_hadir', sum(1 for s in data.get('students', []) if s.get('hadir') == 'YA'))
    
    catatan_val = data.get('catatan', '')
    if not catatan_val:
        catatan_val = '-'
        
    pengawas_val = data.get('nama_pengawas') or data.get('pengawas') or '.............................................'
    if not pengawas_val or not pengawas_val.strip():
        pengawas_val = '.............................................'

    ctx = {
        'jenis_kegiatan': data.get('jenis_kegiatan', 'ONLINE'),
        'hari': data.get('hari', ''),
        'tanggal': data.get('tanggal', ''),
        'waktu': data.get('waktu', f"{data.get('jam_mulai', '09.00')} - {data.get('jam_selesai', '12.00')}"),
        'jumlah_peserta': total_p,
        'jumlah_hadir': hadir_p,
        'catatan': catatan_val,
        'nama_pengawas': pengawas_val
    }
    tpl.render(ctx)
    tpl.save(output_path)
    return output_path

def generate_daftar_hadir_docx(data, output_path, logo_path=None):
    """
    Generates Daftar Hadir Word Document (.docx) using reverensi template.
    Guarantees fixed 28 table rows and dots signature for Pengawas.
    """
    tpl = DocxTemplate(TEMPLATE_DH_PATH)
    
    students_data = data.get('students', [])
    peserta_list = []
    for idx, s in enumerate(students_data):
        peserta_list.append({
            'no': idx + 1,
            'nama': s.get('nama', ''),
            'npm': s.get('npm', ''),
            'hadir': s.get('hadir', '')
        })

    # Always pad table up to at least 28 data rows
    TARGET_ROWS = 28
    while len(peserta_list) < TARGET_ROWS:
        idx = len(peserta_list)
        peserta_list.append({
            'no': idx + 1,
            'nama': '',
            'npm': '',
            'hadir': ''
        })

    pengawas_val = data.get('nama_pengawas') or data.get('pengawas') or '.............................................'
    if not pengawas_val or not pengawas_val.strip():
        pengawas_val = '.............................................'

    ctx = {
        'hari': data.get('hari', ''),
        'tanggal': data.get('tanggal', ''),
        'jam_mulai': data.get('jam_mulai', '09.00'),
        'jam_selesai': data.get('jam_selesai', '12.00'),
        'nama_pengawas': pengawas_val,
        'peserta': peserta_list
    }
    tpl.render(ctx)
    tpl.save(output_path)
    return output_path

def generate_paket_lengkap_bapu_docx(data, output_path, logo_path=None):
    """
    Generates combined BAPU package (Berita Acara + Daftar Hadir) using reverensi templates.
    Strips trailing empty paragraphs to prevent empty page 2.
    """
    ba_temp = output_path + ".ba.docx"
    dh_temp = output_path + ".dh.docx"
    
    generate_berita_acara_docx(data, ba_temp)
    generate_daftar_hadir_docx(data, dh_temp)
    
    doc_master = Document(ba_temp)
    doc_slave = Document(dh_temp)
    
    # Strip trailing empty paragraphs in doc_master to prevent empty page 2!
    while len(doc_master.paragraphs) > 0 and not doc_master.paragraphs[-1].text.strip():
        p_last = doc_master.paragraphs[-1]
        p_last._element.getparent().remove(p_last._element)
        
    doc_master.add_page_break()
    
    for elem in doc_slave.element.body:
        if not elem.tag.endswith('sectPr'):
            doc_master.element.body.append(elem)
            
    doc_master.save(output_path)
    
    if os.path.exists(ba_temp): os.remove(ba_temp)
    if os.path.exists(dh_temp): os.remove(dh_temp)
    
    return output_path
